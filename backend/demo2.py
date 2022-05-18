"""
    demo简易文档纠错演示程序
    读取word文档内容进行分句与纠错后以word形式保存结果
"""

from docx import Document
from docx.shared import Inches
import re
from paddlenlp.transformers import ErnieTokenizer
from paddlenlp.data import Vocab
from predict import Predictor

# 精细的中文分句
def cut_sent(para):
    para = re.sub('([。！？\?])([^”’])', r"\1\n\2", para)  # 单字符断句符
    para = re.sub('(\.{6})([^”’])', r"\1\n\2", para)  # 英文省略号
    para = re.sub('(\…{2})([^”’])', r"\1\n\2", para)  # 中文省略号
    para = re.sub('([。！？\?][”’])([^，。！？\?])', r'\1\n\2', para)
    # 如果双引号前有终止符，那么双引号才是句子的终点，把分句符\n放到双引号后，注意前面的几句都小心保留了双引号
    para = para.rstrip()  # 段尾如果有多余的\n就去掉它
    # 很多规则中会考虑分号;，但是这里我把它忽略不计，破折号、英文双引号等同样忽略，需要的再做些简单调整即可。
    return para.split("\n")

# 读取word文档内容
def get_paragraphs_text(path):
    """
    获取所有段落的文本
    :param path: word路径
    :return: list类型，如：
        ['Test', 'hello world', ...]
    """
    document = Document(path) 
    # 提取正文文本
    paragraphs_text = ""
    for paragraph in document.paragraphs:
        # 拼接一个list,包括段落的结构和内容
        paragraphs_text += paragraph.text + "\n"
    return paragraphs_text

if __name__ == "__main__":
    # 读取要进行文本纠错的word文档
    text = get_paragraphs_text('./resource/test.docx')
    # 对word文档内容进行分句处理
    text = cut_sent(text)
    # 2.加载训练好的文本纠错模型去完成纠错任务
    tokenizer = ErnieTokenizer.from_pretrained("ernie-1.0")
    pinyin_vocab = Vocab.load_vocabulary("./best_model/pinyin_vocab.txt", unk_token='[UNK]', pad_token='[PAD]')
    # 配置加载模型参数地址
    predictor = Predictor('./best_model/static_graph_params.pdmodel', './best_model/static_graph_params.pdiparams', 
                          'cpu', 128, tokenizer, pinyin_vocab)
    # 文本纠错
    results = predictor.predict(text, batch_size=2)
    # 输出纠错后结果
    print("******************************")
    print("文档纠错结果：")
    print("\n".join(results))
    # 文档结果保存
    new_document = Document()
    for idx, item in enumerate(results):
        if item is not '': 
            new_document.add_paragraph(item)
    # 结果保存到新文档    
    new_document.save('./resource/csc-result.docx')
    print("******************************")
    print("纠错结果保存成功！保存路径为：./resource/csc-result.docx")





